import os
import json
import httpx
import asyncio
import time
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status, BackgroundTasks
from sqlalchemy.orm import Session
from pydantic import BaseModel
from app.database import get_db
from app.models.user import User
from app.models.role import Role
from app.models.document import PublicDocument, PersonalDocument, DocumentPermission
from app.services.file_storage import file_storage_service
from app.dependencies import get_current_user
from app.utils.logger import logger
from app.schemas.document import (
    PublicDocumentResponse, 
    PersonalDocumentResponse,
    DocumentUploadResponse,
    DocumentListResponse,
    DocumentDetailResponse,
    FolderUploadResponse,
    FolderUploadResult
)

router = APIRouter()


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    type: str = Form(...),  # "public" 或 "personal"
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    permissions: Optional[str] = Form(None),  # JSON字符串，包含角色ID列表
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)):
    """
    上传文档
    """
    logger.info(f"开始文档上传: 用户ID={current_user.id}, 文件名={file.filename}, 文档类型={type}")
    
    try:
        # 验证文件类型
        allowed_extensions = ['.txt', '.md', '.json', '.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg']
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in allowed_extensions:
            logger.warning(f"不支持的文件类型: {file_extension} (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"不支持的文件类型: {file_extension}。支持的类型: {', '.join(allowed_extensions)}"
            )
        
        # 验证文档类型
        if type not in ["public", "personal"]:
            logger.warning(f"无效的文档类型: {type} (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文档类型必须是 'public' 或 'personal'"
            )
        
        # 如果是公共文档，验证权限
        if type == "public" and not permissions:
            logger.warning(f"公共文档未指定权限 (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="公共文档必须指定权限"
            )
        
        # 解析权限
        permission_role_ids = []
        if permissions:
            try:
                permission_role_ids = json.loads(permissions)
                if not isinstance(permission_role_ids, list):
                    raise ValueError("权限必须是角色ID列表")
                logger.debug(f"解析权限成功: {permission_role_ids} (用户ID={current_user.id})")
            except (json.JSONDecodeError, ValueError) as e:
                logger.error(f"权限解析失败: {str(e)} (用户ID={current_user.id})")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"权限格式错误: {str(e)}"
                )
        
        # 保存文件
        logger.debug(f"开始保存文件: {file.filename} (用户ID={current_user.id})")
        file_hash, file_path, file_size, file_type, minio_filename, file_url = await file_storage_service.save_file(file, type)
        logger.info(f"文件保存成功: 哈希={file_hash}, 路径={file_path}, 大小={file_size}字节, MinIO文件名={minio_filename}, MinIO URL={file_url} (用户ID={current_user.id})")
        
        # 检查文件是否已存在
        if type == "public":
            existing_doc = db.query(PublicDocument).filter(PublicDocument.file_hash == file_hash).first()
            if existing_doc:
                logger.info(f"公共文档已存在: {file.filename} (哈希: {file_hash})")
                return DocumentUploadResponse(
                    success=True,
                    message="文档已存在",
                    document_id=existing_doc.id,
                    document_type="public",
                    filename=existing_doc.filename,
                    file_path=existing_doc.file_path,
                    file_size=existing_doc.file_size,
                    permissions={"roles": permission_role_ids} if permission_role_ids else None
                )
        else:  # personal
            existing_doc = db.query(PersonalDocument).filter(
                PersonalDocument.file_hash == file_hash,
                PersonalDocument.owner_id == current_user.id
            ).first()
            if existing_doc:
                logger.info(f"个人文档已存在: {file.filename} (哈希: {file_hash})")
                return DocumentUploadResponse(
                    success=True,
                    message="文档已存在",
                    document_id=existing_doc.id,
                    document_type="personal",
                    filename=existing_doc.filename,
                    file_path=existing_doc.file_path,
                    file_size=existing_doc.file_size,
                    permissions=None
                )
        
        # 创建数据库记录
        if type == "public":
            # 创建公共文档记录
            new_doc = PublicDocument(
                filename=file.filename,
                file_path=file_path,
                file_size=file_size,
                file_type=file_type,
                file_hash=file_hash,
                minio_filename=minio_filename,
                file_url=file_url,
                uploader_id=current_user.id,
                title=title,
                description=description,
                is_active=True,
                is_processed=False  # RAG处理默认为false
            )
            db.add(new_doc)
            db.flush()  # 获取新创建的文档ID
            
            # 添加权限记录
            for role_id in permission_role_ids:
                permission = DocumentPermission(
                    document_id=new_doc.id,
                    role_id=role_id,
                    granted_by=current_user.id
                )
                db.add(permission)
            
            db.commit()
            
            logger.info(f"公共文档上传成功: {file.filename} (ID: {new_doc.id}, 用户ID={current_user.id}, 权限角色: {permission_role_ids})")
            
            return DocumentUploadResponse(
                success=True,
                message="公共文档上传成功",
                document_id=new_doc.id,
                document_type="public",
                filename=new_doc.filename,
                file_path=new_doc.file_path,
                file_size=new_doc.file_size,
                permissions={"roles": permission_role_ids}
            )
            
        else:  # personal
            # 创建个人文档记录
            new_doc = PersonalDocument(
                filename=file.filename,
                file_path=file_path,
                file_size=file_size,
                file_type=file_type,
                file_hash=file_hash,
                minio_filename=minio_filename,
                file_url=file_url,
                owner_id=current_user.id,
                title=title,
                description=description,
                is_active=True,
                is_processed=False  # RAG处理默认为false
            )
            db.add(new_doc)
            db.commit()
            
            logger.info(f"个人文档上传成功: {file.filename} (ID: {new_doc.id}, 用户ID={current_user.id})")
            
            return DocumentUploadResponse(
                success=True,
                message="个人文档上传成功",
                document_id=new_doc.id,
                document_type="personal",
                filename=new_doc.filename,
                file_path=new_doc.file_path,
                file_size=new_doc.file_size,
                permissions=None
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"文档上传失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"文档上传失败: {str(e)}"
        )


@router.post("/upload-folder", response_model=FolderUploadResponse)
async def upload_folder(
    files: List[UploadFile] = File(...),
    type: str = Form(...),  # "public" 或 "personal"
    permissions: Optional[str] = Form(None),  # JSON字符串，包含角色ID列表
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)):
    """
    上传文件夹中的多个文档
    """
    logger.info(f"开始文件夹上传: 用户ID={current_user.id}, 文件数量={len(files)}, 文档类型={type}")
    
    results = []
    
    # 验证文档类型
    if type not in ["public", "personal"]:
        logger.warning(f"无效的文档类型: {type} (用户ID={current_user.id})")
        return FolderUploadResponse(
            success=False,
            message=f"无效的文档类型: {type}",
            results=[]
        )
    
    # 如果是公共文档，验证权限
    if type == "public" and not permissions:
        logger.warning(f"公共文档未指定权限 (用户ID={current_user.id})")
        return FolderUploadResponse(
            success=False,
            message="公共文档必须指定权限",
            results=[]
        )
    
    # 解析权限
    permission_role_ids = []
    if permissions:
        try:
            permission_role_ids = json.loads(permissions)
            if not isinstance(permission_role_ids, list):
                raise ValueError("权限必须是角色ID列表")
            logger.debug(f"解析权限成功: {permission_role_ids} (用户ID={current_user.id})")
        except (json.JSONDecodeError, ValueError) as e:
            logger.error(f"权限解析失败: {str(e)} (用户ID={current_user.id})")
            return FolderUploadResponse(
                success=False,
                message=f"权限格式错误: {str(e)}",
                results=[]
            )
    
    # 验证文件类型
    allowed_extensions = ['.txt', '.md', '.json', '.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg']
    
    # 逐个处理文件
    for file in files:
        try:
            # 验证文件类型
            file_extension = os.path.splitext(file.filename)[1].lower()
            if file_extension not in allowed_extensions:
                logger.warning(f"跳过不支持的文件类型: {file.filename} (用户ID={current_user.id})")
                results.append(FolderUploadResult(
                    filename=file.filename,
                    success=False,
                    message=f"不支持的文件类型: {file_extension}"
                ))
                continue
            
            # 保存文件
            logger.debug(f"开始保存文件: {file.filename} (用户ID={current_user.id})")
            file_hash, file_path, file_size, file_type, minio_filename, file_url = await file_storage_service.save_file(file, type)
            logger.info(f"文件保存成功: 哈希={file_hash}, 路径={file_path}, 大小={file_size}字节 (用户ID={current_user.id})")
            
            # 检查文件是否已存在
            if type == "public":
                existing_doc = db.query(PublicDocument).filter(PublicDocument.file_hash == file_hash).first()
                if existing_doc:
                    logger.info(f"公共文档已存在: {file.filename} (哈希: {file_hash})")
                    results.append(FolderUploadResult(
                        filename=file.filename,
                        success=True,
                        message="文档已存在",
                        document_id=existing_doc.id,
                        document_type="public"
                    ))
                    continue
            else:  # personal
                existing_doc = db.query(PersonalDocument).filter(
                    PersonalDocument.file_hash == file_hash,
                    PersonalDocument.owner_id == current_user.id
                ).first()
                if existing_doc:
                    logger.info(f"个人文档已存在: {file.filename} (哈希: {file_hash})")
                    results.append(FolderUploadResult(
                        filename=file.filename,
                        success=True,
                        message="文档已存在",
                        document_id=existing_doc.id,
                        document_type="personal"
                    ))
                    continue
            
            # 创建数据库记录
            if type == "public":
                # 创建公共文档记录
                new_doc = PublicDocument(
                    filename=file.filename,
                    file_path=file_path,
                    file_size=file_size,
                    file_type=file_type,
                    file_hash=file_hash,
                    minio_filename=minio_filename,
                    file_url=file_url,
                    uploader_id=current_user.id,
                    title=file.filename,
                    description=None,
                    is_active=True,
                    is_processed=False
                )
                db.add(new_doc)
                db.flush()
                
                # 添加权限记录
                for role_id in permission_role_ids:
                    permission = DocumentPermission(
                        document_id=new_doc.id,
                        role_id=role_id,
                        granted_by=current_user.id
                    )
                    db.add(permission)
                
                db.commit()
                
                logger.info(f"公共文档上传成功: {file.filename} (ID: {new_doc.id}, 用户ID={current_user.id})")
                
                results.append(FolderUploadResult(
                    filename=file.filename,
                    success=True,
                    message="上传成功",
                    document_id=new_doc.id,
                    document_type="public"
                ))
                
            else:  # personal
                # 创建个人文档记录
                new_doc = PersonalDocument(
                    filename=file.filename,
                    file_path=file_path,
                    file_size=file_size,
                    file_type=file_type,
                    file_hash=file_hash,
                    minio_filename=minio_filename,
                    file_url=file_url,
                    owner_id=current_user.id,
                    title=file.filename,
                    description=None,
                    is_active=True,
                    is_processed=False
                )
                db.add(new_doc)
                db.commit()
                
                logger.info(f"个人文档上传成功: {file.filename} (ID: {new_doc.id}, 用户ID={current_user.id})")
                
                results.append(FolderUploadResult(
                    filename=file.filename,
                    success=True,
                    message="上传成功",
                    document_id=new_doc.id,
                    document_type="personal"
                ))
                
        except Exception as e:
            logger.error(f"文件上传失败: {file.filename}, 错误: {str(e)} (用户ID={current_user.id})")
            results.append(FolderUploadResult(
                filename=file.filename,
                success=False,
                message=f"上传失败: {str(e)}"
            ))
    
    # 统计成功和失败的数量
    success_count = sum(1 for r in results if r.success)
    failed_count = len(results) - success_count
    
    logger.info(f"文件夹上传完成: 总数={len(results)}, 成功={success_count}, 失败={failed_count} (用户ID={current_user.id})")
    
    return FolderUploadResponse(
        success=True,
        message=f"上传完成: 成功 {success_count} 个, 失败 {failed_count} 个",
        results=results
    )


@router.get("/public", response_model=List[PublicDocumentResponse])
def get_public_documents(
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取公共文档列表
    """
    logger.info(f"获取公共文档列表: 用户ID={current_user.id}, skip={skip}, limit={limit}")
    
    try:
        # 这里需要根据用户权限过滤文档
        # 暂时返回所有公共文档，后续需要实现权限过滤
        documents = db.query(PublicDocument).filter(
            PublicDocument.is_active == True
        ).offset(skip).limit(limit).all()
        
        # 为每个文档添加权限信息
        documents_with_permissions = []
        for doc in documents:
            # 获取文档的权限角色ID列表
            permissions = db.query(DocumentPermission.role_id).filter(
                DocumentPermission.document_id == doc.id
            ).all()
            permission_role_ids = [perm.role_id for perm in permissions]
            
            # 创建包含权限信息的文档对象
            doc_dict = {
                'id': doc.id,
                'filename': doc.filename,
                'title': doc.title,
                'description': doc.description,
                'file_path': doc.file_path,
                'file_size': doc.file_size,
                'file_type': doc.file_type,
                'file_hash': doc.file_hash,
                'uploader_id': doc.uploader_id,
                'upload_time': doc.upload_time,
                'is_active': doc.is_active,
                'is_processed': doc.is_processed,
                'created_at': doc.created_at,
                'updated_at': doc.updated_at,
                'permissions': permission_role_ids
            }
            documents_with_permissions.append(doc_dict)
        
        logger.info(f"成功获取公共文档列表: 用户ID={current_user.id}, 文档数量={len(documents_with_permissions)}")
        return documents_with_permissions
    except Exception as e:
        logger.error(f"获取公共文档列表失败: {str(e)} (用户ID={current_user.id})")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取公共文档列表失败: {str(e)}"
        )


@router.get("/personal", response_model=List[PersonalDocumentResponse])
def get_personal_documents(
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取个人文档列表
    """
    logger.info(f"获取个人文档列表: 用户ID={current_user.id}, skip={skip}, limit={limit}")
    
    try:
        documents = db.query(PersonalDocument).filter(
            PersonalDocument.owner_id == current_user.id,
            PersonalDocument.is_active == True
        ).offset(skip).limit(limit).all()
        
        # 为每个文档添加权限信息（个人文档默认只有自己可见）
        documents_with_permissions = []
        for doc in documents:
            # 个人文档没有权限控制，但为了前端一致性，添加空权限列表
            doc_dict = {
                'id': doc.id,
                'filename': doc.filename,
                'title': doc.title,
                'description': doc.description,
                'file_path': doc.file_path,
                'file_size': doc.file_size,
                'file_type': doc.file_type,
                'file_hash': doc.file_hash,
                'owner_id': doc.owner_id,
                'upload_time': doc.upload_time,
                'is_active': doc.is_active,
                'is_processed': doc.is_processed,
                'created_at': doc.created_at,
                'updated_at': doc.updated_at,
                'permissions': []
            }
            documents_with_permissions.append(doc_dict)
        
        logger.info(f"成功获取个人文档列表: 用户ID={current_user.id}, 文档数量={len(documents_with_permissions)}")
        return documents_with_permissions
    except Exception as e:
        logger.error(f"获取个人文档列表失败: {str(e)} (用户ID={current_user.id})")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取个人文档列表失败: {str(e)}"
        )


@router.get("/{document_type}/{document_id}/content")
def get_document_content(
    document_type: str,
    document_id: int,
    page: Optional[int] = None,  # 新增：分页参数
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取文档内容
    """
    logger.info(f"获取文档内容: 用户ID={current_user.id}, 文档类型={document_type}, 文档ID={document_id}, 页码={page}")
    
    try:
        # 验证文档类型
        if document_type not in ["public", "personal"]:
            logger.warning(f"无效的文档类型: {document_type} (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文档类型必须是 'public' 或 'personal'"
            )
        
        # 获取文档
        if document_type == "public":
            document = db.query(PublicDocument).filter(
                PublicDocument.id == document_id,
                PublicDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在"
                )
                
            # TODO: 检查用户是否有权限访问此文档
            # 暂时允许所有用户访问公共文档
            
        else:  # personal
            document = db.query(PersonalDocument).filter(
                PersonalDocument.id == document_id,
                PersonalDocument.owner_id == current_user.id,
                PersonalDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在或没有权限访问"
                )
        
        # 读取文件内容
        content = ""
        file_extension = Path(document.file_path).suffix.lower()
        
        try:
            # 使用file_storage_service获取完整路径
            full_file_path = file_storage_service.get_file_path(document.file_path)
            
            # 根据文件类型处理
            if file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml']:
                # 文本文件，直接读取
                try:
                    with open(full_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    logger.info(f"成功读取文本内容: 文档ID={document_id}, 内容长度={len(content)}字符")
                except UnicodeDecodeError:
                    # 如果UTF-8解码失败，尝试其他编码
                    with open(full_file_path, 'r', encoding='gbk') as f:
                        content = f.read()
                    logger.info(f"成功读取文本内容(GBK编码): 文档ID={document_id}, 内容长度={len(content)}字符")
            elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg']:
                # 图片文件，返回图片信息而不是二进制内容
                import base64
                with open(full_file_path, 'rb') as f:
                    image_data = f.read()
                
                # 将图片转换为base64编码
                base64_image = base64.b64encode(image_data).decode('utf-8')
                
                # 获取图片尺寸信息
                try:
                    if file_extension in ['.jpg', '.jpeg']:
                        from PIL import Image
                        img = Image.open(full_file_path)
                        width, height = img.size
                        content = f"[图片信息]\n格式: {file_extension}\n尺寸: {width}x{height}\n大小: {document.file_size}字节\nBase64编码: {base64_image}"
                    else:
                        content = f"[图片信息]\n格式: {file_extension}\n大小: {document.file_size}字节\nBase64编码: {base64_image}"
                except ImportError:
                    content = f"[图片信息]\n格式: {file_extension}\n大小: {document.file_size}字节\nBase64编码: {base64_image}"
                except Exception as e:
                    content = f"[图片信息]\n格式: {file_extension}\n大小: {document.file_size}字节\nBase64编码: {base64_image}\n错误: 无法获取图片尺寸 - {str(e)}"
                
                logger.info(f"成功读取图片内容: 文档ID={document_id}, 格式={file_extension}, 大小={document.file_size}字节")
            elif file_extension in ['.docx', '.doc']:
                # Word文档，返回Base64编码的原始数据
                import base64
                try:
                    with open(full_file_path, 'rb') as file:
                        docx_data = file.read()
                    
                    # 将Word文档转换为base64编码
                    base64_docx = base64.b64encode(docx_data).decode('utf-8')
                    
                    # 尝试提取文本用于搜索和预览
                    try:
                        from docx import Document
                        doc = Document(full_file_path)
                        text_content = '\n'.join([para.text for para in doc.paragraphs])
                        
                        # 获取文档基本信息
                        core_props = doc.core_properties
                        author = core_props.author or "未知"
                        title = core_props.title or "无标题"
                        created = core_props.created or "未知"
                        
                        content = f"[Word文档]\n格式: {file_extension}\n大小: {document.file_size}字节\n作者: {author}\n标题: {title}\n创建时间: {created}\nBase64编码: {base64_docx}\n[提取文本]\n{text_content}"
                    except Exception as e:
                        logger.warning(f"提取Word文本失败，但仍返回原始文档数据: 文档ID={document_id}, 错误={str(e)}")
                        content = f"[Word文档]\n格式: {file_extension}\n大小: {document.file_size}字节\nBase64编码: {base64_docx}"
                    
                    logger.info(f"成功读取Word文档: 文档ID={document_id}, 大小={document.file_size}字节")
                except Exception as e:
                    logger.error(f"读取Word文档失败: 文档ID={document_id}, 错误={str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"读取Word文档失败: {str(e)}"
                    )
            elif file_extension == '.pdf':
                # PDF文档 - 优化：支持分页加载
                import base64
                try:
                    # 获取PDF基本信息
                    import PyPDF2
                    reader = PyPDF2.PdfReader(full_file_path)
                    total_pages = len(reader.pages)
                    
                    # 如果指定了页码，只返回该页的内容
                    if page is not None:
                        if page < 1 or page > total_pages:
                            raise HTTPException(
                                status_code=status.HTTP_400_BAD_REQUEST,
                                detail=f"页码超出范围: 有效范围 1-{total_pages}"
                            )
                        
                        # 提取单页文本
                        page_text = reader.pages[page-1].extract_text()
                        
                        # 将单页转换为base64（仅该页的PDF数据）
                        import io
                        from PyPDF2 import PdfWriter
                        
                        writer = PdfWriter()
                        writer.add_page(reader.pages[page-1])
                        
                        pdf_buffer = io.BytesIO()
                        writer.write(pdf_buffer)
                        pdf_buffer.seek(0)
                        
                        base64_page_pdf = base64.b64encode(pdf_buffer.read()).decode('utf-8')
                        
                        content = f"[PDF文档-单页]\n格式: PDF\n大小: {document.file_size}字节\n总页数: {total_pages}\n当前页: {page}\nBase64编码: {base64_page_pdf}\n[提取文本]\n{page_text}"
                        
                        logger.info(f"成功读取PDF单页内容: 文档ID={document_id}, 页码={page}, 页大小={len(base64_page_pdf)}字符")
                    
                    else:
                        # 默认行为：返回完整PDF（保持向后兼容）
                        with open(full_file_path, 'rb') as file:
                            pdf_data = file.read()
                        
                        base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
                        
                        # 提取所有页面的文本
                        text_content = ""
                        for page_num, pdf_page in enumerate(reader.pages, 1):
                            text_content += f"第{page_num}页:\n{pdf_page.extract_text()}\n\n"
                        
                        content = f"[PDF文档-完整]\n格式: PDF\n大小: {document.file_size}字节\n页数: {total_pages}\nBase64编码: {base64_pdf}\n[提取文本]\n{text_content}"
                        
                        logger.info(f"成功读取完整PDF内容: 文档ID={document_id}, 大小={document.file_size}字节")
                    
                except Exception as e:
                    logger.error(f"读取PDF文档失败: 文档ID={document_id}, 错误={str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"读取PDF文档失败: {str(e)}"
                    )
            else:
                # 不支持的文件类型
                logger.error(f"不支持的文件类型: {file_extension} (文档ID={document_id})")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"不支持的文件类型: {file_extension}，支持的格式包括：.txt, .md, .docx, .pdf, .png, .jpg, .jpeg, .gif, .bmp, .svg等"
                )
                
        except Exception as e:
            logger.error(f"读取文档失败: 文档ID={document_id}, 错误={str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"读取文档失败: {str(e)}"
            )
        
        # 按行分割内容
        lines = content.split('\n')
        
        return {
            "document_id": document_id,
            "document_type": document_type,
            "filename": document.filename,
            "file_size": document.file_size,
            "file_type": document.file_type,
            "content": content,
            "lines": lines,
            "line_count": len(lines),
            "page": page,  # 返回当前页码（如果指定）
            "total_pages": total_pages if file_extension == '.pdf' else None  # 返回总页数（仅PDF）
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文档内容失败: {str(e)} (用户ID={current_user.id})")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取文档内容失败: {str(e)}"
        )


@router.get("/{document_type}/{document_id}", response_model=DocumentDetailResponse)
def get_document_detail(
    document_type: str,
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取文档详细信息
    """
    logger.info(f"获取文档详细信息: 用户ID={current_user.id}, 文档类型={document_type}, 文档ID={document_id}")
    
    try:
        # 验证文档类型
        if document_type not in ["public", "personal"]:
            logger.warning(f"无效的文档类型: {document_type} (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文档类型必须是 'public' 或 'personal'"
            )
        
        # 获取文档
        if document_type == "public":
            document = db.query(PublicDocument).filter(
                PublicDocument.id == document_id,
                PublicDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在"
                )
            
            # 获取文档的权限角色ID列表
            permissions = db.query(DocumentPermission.role_id).filter(
                DocumentPermission.document_id == document_id
            ).all()
            permission_role_ids = [perm.role_id for perm in permissions]
            
            # 获取上传者信息
            uploader = db.query(User).filter(User.id == document.uploader_id).first()
            
            # 详细日志输出
            logger.info("================================================================================")
            logger.info("【文档详细信息检索】")
            logger.info(f"文档ID: {document.id}")
            logger.info(f"文件名: {document.filename}")
            logger.info(f"文件路径: {document.file_path}")
            logger.info(f"文件大小: {document.file_size} 字节")
            logger.info(f"文件类型: {document.file_type}")
            logger.info(f"文件哈希: {document.file_hash}")
            logger.info(f"MinIO文件名: {document.minio_filename or '无'}")
            logger.info(f"文件URL: {document.file_url or '无'}")
            logger.info(f"文档标题: {document.title or '无'}")
            logger.info(f"文档描述: {document.description or '无'}")
            logger.info(f"文档类型: public")
            logger.info(f"是否激活: {'是' if document.is_active else '否'}")
            logger.info(f"是否已处理: {'是' if document.is_processed else '否'}")
            logger.info(f"上传时间: {document.upload_time}")
            logger.info(f"创建时间: {document.created_at}")
            logger.info(f"更新时间: {document.updated_at}")
            logger.info(f"上传者信息: {{'user_id': {document.uploader_id}, 'username': {uploader.username if uploader else 'N/A'}, 'email': {uploader.email if uploader else 'N/A'}, 'role_ids': {uploader.get_role_ids() if uploader else 'N/A'}}}")
            logger.info(f"所属部门数量: {len(permission_role_ids)}")
            for idx, role_id in enumerate(permission_role_ids, 1):
                role = db.query(Role).filter(Role.id == role_id).first()
                logger.info(f"  部门{idx}:")
                logger.info(f"    - 角色ID: {role_id}")
                if role:
                    logger.info(f"    - 角色代码: {role.role_code}")
                    logger.info(f"    - 角色名称: {role.role_name}")
                    logger.info(f"    - 角色描述: {role.description}")
            logger.info("================================================================================")
            
            # 构建响应
            detail = {
                "id": document.id,
                "filename": document.filename,
                "file_path": document.file_path,
                "file_size": document.file_size,
                "file_type": document.file_type,
                "file_hash": document.file_hash,
                "minio_filename": document.minio_filename,
                "file_url": document.file_url,
                "title": document.title,
                "description": document.description,
                "is_active": document.is_active,
                "is_processed": document.is_processed,
                "upload_time": document.upload_time,
                "created_at": document.created_at,
                "updated_at": document.updated_at,
                "document_type": "public",
                "uploader_id": document.uploader_id,
                "owner_id": None,
                "permissions": permission_role_ids
            }
            
            logger.info(f"成功获取公共文档详细信息: 文档ID={document_id}, 文件名={document.filename}")
            
        else:  # personal
            document = db.query(PersonalDocument).filter(
                PersonalDocument.id == document_id,
                PersonalDocument.owner_id == current_user.id,
                PersonalDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在或没有权限访问"
                )
            
            # 获取所有者信息
            owner = db.query(User).filter(User.id == document.owner_id).first()
            
            # 详细日志输出
            logger.info("================================================================================")
            logger.info("【文档详细信息检索】")
            logger.info(f"文档ID: {document.id}")
            logger.info(f"文件名: {document.filename}")
            logger.info(f"文件路径: {document.file_path}")
            logger.info(f"文件大小: {document.file_size} 字节")
            logger.info(f"文件类型: {document.file_type}")
            logger.info(f"文件哈希: {document.file_hash}")
            logger.info(f"MinIO文件名: {document.minio_filename or '无'}")
            logger.info(f"文件URL: {document.file_url or '无'}")
            logger.info(f"文档标题: {document.title or '无'}")
            logger.info(f"文档描述: {document.description or '无'}")
            logger.info(f"文档类型: personal")
            logger.info(f"是否激活: {'是' if document.is_active else '否'}")
            logger.info(f"是否已处理: {'是' if document.is_processed else '否'}")
            logger.info(f"上传时间: {document.upload_time}")
            logger.info(f"创建时间: {document.created_at}")
            logger.info(f"更新时间: {document.updated_at}")
            logger.info(f"所有者信息: {{'user_id': {document.owner_id}, 'username': {owner.username if owner else 'N/A'}, 'email': {owner.email if owner else 'N/A'}, 'role_ids': {owner.get_role_ids() if owner else 'N/A'}}}")
            logger.info("================================================================================")
            
            # 构建响应
            detail = {
                "id": document.id,
                "filename": document.filename,
                "file_path": document.file_path,
                "file_size": document.file_size,
                "file_type": document.file_type,
                "file_hash": document.file_hash,
                "minio_filename": document.minio_filename,
                "file_url": document.file_url,
                "title": document.title,
                "description": document.description,
                "is_active": document.is_active,
                "is_processed": document.is_processed,
                "upload_time": document.upload_time,
                "created_at": document.created_at,
                "updated_at": document.updated_at,
                "document_type": "personal",
                "uploader_id": None,
                "owner_id": document.owner_id,
                "permissions": None
            }
            
            logger.info(f"成功获取个人文档详细信息: 文档ID={document_id}, 文件名={document.filename}")
        
        return detail
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文档详细信息失败: {str(e)} (用户ID={current_user.id})")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取文档详细信息失败: {str(e)}"
        )


@router.delete("/public/{document_id}")
def delete_public_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    删除公共文档
    """
    try:
        document = db.query(PublicDocument).filter(
            PublicDocument.id == document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="文档不存在"
            )
        
        # 检查权限：只有上传者或管理员可以删除
        if document.uploader_id != current_user.id and "ADMIN" not in current_user.get_role_ids():
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="没有权限删除此文档"
            )
        
        # 删除文件
        file_storage_service.delete_file(document.file_path)
        
        # 删除数据库记录（级联删除权限）
        db.delete(document)
        db.commit()
        
        logger.info(f"公共文档已删除: {document.filename} (ID: {document_id})")
        
        return {"message": "公共文档删除成功"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除公共文档失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"删除公共文档失败: {str(e)}"
        )


@router.delete("/personal/{document_id}")
def delete_personal_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    删除个人文档
    """
    try:
        document = db.query(PersonalDocument).filter(
            PersonalDocument.id == document_id
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="文档不存在"
            )
        
        # 检查权限：只有所有者可以删除
        if document.owner_id != current_user.id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="没有权限删除此文档"
            )
        
        # 删除文件
        file_storage_service.delete_file(document.file_path)
        
        # 删除数据库记录
        db.delete(document)
        db.commit()
        
        logger.info(f"个人文档已删除: {document.filename} (ID: {document_id})")
        
        return {"message": "个人文档删除成功"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除个人文档失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"删除个人文档失败: {str(e)}"
        )


async def poll_document_status(track_id: str, knowledge_name: str, document_id: int, document_type: str):
    """
    后台轮询文档解析状态
    """
    status_url = "http://10.168.27.191:8888/document/status"
    max_attempts = 20
    poll_interval = 30
    
    try:
        for attempt in range(max_attempts):
            try:
                await asyncio.sleep(poll_interval)
                
                async with httpx.AsyncClient() as client:
                    status_response = await client.get(
                        status_url,
                        params={
                            "track_id": track_id,
                            "knowledge_name": knowledge_name
                        }
                    )
                    status_response.raise_for_status()
                    status_result = status_response.json()
                    logger.info(f"文档解析状态 (第{attempt + 1}次查询): {status_result}")
                    
                    status = status_result.get("status")
                    
                    if status == "completed":
                        db = next(get_db())
                        try:
                            if document_type == "public":
                                document = db.query(PublicDocument).filter(PublicDocument.id == document_id).first()
                            else:
                                document = db.query(PersonalDocument).filter(PersonalDocument.id == document_id).first()
                            
                            if document:
                                document.is_processed = True
                                db.commit()
                                logger.info(f"文档解析完成，已更新处理状态: {document.filename} (ID: {document_id})")
                            else:
                                logger.error(f"文档不存在，无法更新处理状态: document_id={document_id}, document_type={document_type}")
                        finally:
                            db.close()
                        break
                    elif status == "failed":
                        error_message = status_result.get("error_message") or "文档解析失败"
                        logger.error(f"文档解析失败: {error_message}")
                        break
                    else:
                        logger.info(f"文档解析中 (状态: {status})，等待下一次查询...")
                        continue
            except Exception as e:
                logger.warning(f"第{attempt + 1}次查询文档解析状态失败: {str(e)}")
                if attempt == max_attempts - 1:
                    logger.error(f"查询文档解析状态超时: 已尝试{max_attempts}次")
    except Exception as e:
        logger.error(f"轮询文档状态时发生异常: {str(e)}")


@router.post("/{document_type}/{document_id}/process")
async def process_document(
    background_tasks: BackgroundTasks,
    document_type: str,
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    处理文档（向量化等）
    预留接口，待后续实现
    """
    logger.info(f"处理文档请求: 用户ID={current_user.id}, 文档类型={document_type}, 文档ID={document_id}")
    
    try:
        # 验证文档类型
        if document_type not in ["public", "personal"]:
            logger.warning(f"无效的文档类型: {document_type} (用户ID={current_user.id})")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文档类型必须是 'public' 或 'personal'"
            )
        
        # 获取文档
        if document_type == "public":
            document = db.query(PublicDocument).filter(
                PublicDocument.id == document_id,
                PublicDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在"
                )
            
            # 获取上传者信息
            uploader = db.query(User).filter(User.id == document.uploader_id).first()
            
            # 获取文档的权限角色ID列表
            permissions = db.query(DocumentPermission.role_id).filter(
                DocumentPermission.document_id == document_id
            ).all()
            permission_role_ids = [perm.role_id for perm in permissions]
            
            # 详细日志输出
            logger.info("================================================================================")
            logger.info("【文档信息检索】")
            logger.info(f"文档ID: {document.id}")
            logger.info(f"文件名: {document.filename}")
            logger.info(f"文件路径: {document.file_path}")
            logger.info(f"文件大小: {document.file_size} 字节")
            logger.info(f"文件类型: {document.file_type}")
            logger.info(f"文件哈希: {document.file_hash}")
            logger.info(f"MinIO文件名: {document.minio_filename or '无'}")
            logger.info(f"文件URL: {document.file_url or '无'}")
            logger.info(f"文档标题: {document.title or '无'}")
            logger.info(f"文档描述: {document.description or '无'}")
            logger.info(f"文档类型: public")
            logger.info(f"是否激活: {'是' if document.is_active else '否'}")
            logger.info(f"是否已处理: {'是' if document.is_processed else '否'}")
            logger.info(f"上传时间: {document.upload_time}")
            logger.info(f"创建时间: {document.created_at}")
            logger.info(f"更新时间: {document.updated_at}")
            logger.info(f"上传者信息: {{'user_id': {document.uploader_id}, 'username': {uploader.username if uploader else 'N/A'}, 'email': {uploader.email if uploader else 'N/A'}, 'role_ids': {uploader.get_role_ids() if uploader else 'N/A'}}}")
            logger.info(f"所属部门数量: {len(permission_role_ids)}")
            for idx, role_id in enumerate(permission_role_ids, 1):
                role = db.query(Role).filter(Role.id == role_id).first()
                logger.info(f"  部门{idx}:")
                logger.info(f"    - 角色ID: {role_id}")
                if role:
                    logger.info(f"    - 角色代码: {role.role_code}")
                    logger.info(f"    - 角色名称: {role.role_name}")
                    logger.info(f"    - 角色描述: {role.description}")
            logger.info("================================================================================")
                
        else:  # personal
            document = db.query(PersonalDocument).filter(
                PersonalDocument.id == document_id,
                PersonalDocument.owner_id == current_user.id,
                PersonalDocument.is_active == True
            ).first()
            
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="文档不存在或没有权限访问"
                )
            
            # 获取所有者信息
            owner = db.query(User).filter(User.id == document.owner_id).first()
            
            # 详细日志输出
            logger.info("================================================================================")
            logger.info("【文档信息检索】")
            logger.info(f"文档ID: {document.id}")
            logger.info(f"文件名: {document.filename}")
            logger.info(f"文件路径: {document.file_path}")
            logger.info(f"文件大小: {document.file_size} 字节")
            logger.info(f"文件类型: {document.file_type}")
            logger.info(f"文件哈希: {document.file_hash}")
            logger.info(f"MinIO文件名: {document.minio_filename or '无'}")
            logger.info(f"文件URL: {document.file_url or '无'}")
            logger.info(f"文档标题: {document.title or '无'}")
            logger.info(f"文档描述: {document.description or '无'}")
            logger.info(f"文档类型: personal")
            logger.info(f"是否激活: {'是' if document.is_active else '否'}")
            logger.info(f"是否已处理: {'是' if document.is_processed else '否'}")
            logger.info(f"上传时间: {document.upload_time}")
            logger.info(f"创建时间: {document.created_at}")
            logger.info(f"更新时间: {document.updated_at}")
            logger.info(f"所有者信息: {{'user_id': {document.owner_id}, 'username': {owner.username if owner else 'N/A'}, 'email': {owner.email if owner else 'N/A'}, 'role_ids': {owner.get_role_ids() if owner else 'N/A'}}}")
            logger.info("================================================================================")
        
        # 检查文档是否已处理
        if document.is_processed:
            logger.info(f"文档已处理: {document.filename} (ID: {document_id})")
            return {
                "success": True,
                "message": "文档已处理",
                "document_id": document_id
            }
        
        # 查询现有知识库
        knowledge_base_url = "http://10.168.27.191:8888/knowledges"
        try:
            with httpx.Client() as client:
                response = client.get(knowledge_base_url)
                response.raise_for_status()
                existing_knowledges = response.json()
                logger.info(f"知识库API返回数据: {existing_knowledges}")
                if not isinstance(existing_knowledges, list):
                    logger.error(f"知识库API返回数据格式错误，期望列表，实际类型: {type(existing_knowledges)}")
                    existing_knowledges = []
            logger.info(f"现有知识库数量: {len(existing_knowledges)}")
        except Exception as e:
            logger.error(f"获取知识库列表失败: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"获取知识库列表失败: {str(e)}"
            )
        
        # 检查并创建知识库
        if document_type == "personal":
            # 个人文档：检查是否有knowledge_personal知识库
            knowledge_name = "knowledge_personal"
            knowledge_exists = any(k == knowledge_name for k in existing_knowledges)
            
            if not knowledge_exists:
                logger.info(f"创建个人知识库: {knowledge_name}")
                try:
                    with httpx.Client() as client:
                        create_response = client.post(
                            knowledge_base_url,
                            data={
                                "knowledge_name": knowledge_name,
                                "knowledge_summary": "个人文档知识库"
                            }
                        )
                        create_response.raise_for_status()
                    logger.info(f"个人知识库创建成功: {knowledge_name}")
                except Exception as e:
                    logger.error(f"创建个人知识库失败: {str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"创建个人知识库失败: {str(e)}"
                    )
            else:
                logger.info(f"个人知识库已存在: {knowledge_name}")
        else:
            # 公共文档：检查是否有对应角色代码的知识库
            permissions = db.query(DocumentPermission.role_id).filter(
                DocumentPermission.document_id == document_id
            ).all()
            permission_role_ids = [perm.role_id for perm in permissions]
            
            for role_id in permission_role_ids:
                role = db.query(Role).filter(Role.id == role_id).first()
                if role and role.role_code:
                    knowledge_name = f"knowledge_{role.role_code}"
                    knowledge_exists = any(k == knowledge_name for k in existing_knowledges)
                    
                    if not knowledge_exists:
                        logger.info(f"创建角色知识库: {knowledge_name} (角色代码: {role.role_code})")
                        try:
                            with httpx.Client() as client:
                                create_response = client.post(
                                    knowledge_base_url,
                                    data={
                                        "knowledge_name": knowledge_name,
                                        "knowledge_summary": f"{role.role_name}知识库"
                                    }
                                )
                                create_response.raise_for_status()
                            logger.info(f"角色知识库创建成功: {knowledge_name}")
                        except Exception as e:
                            logger.error(f"创建角色知识库失败: {str(e)}")
                            raise HTTPException(
                                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                                detail=f"创建角色知识库失败: {str(e)}"
                            )
                    else:
                        logger.info(f"角色知识库已存在: {knowledge_name}")
        
        # 确定知识库名称
        if document_type == "personal":
            knowledge_name = "knowledge_personal"
        else:
            # 公共文档使用第一个角色的知识库
            permissions = db.query(DocumentPermission.role_id).filter(
                DocumentPermission.document_id == document_id
            ).all()
            permission_role_ids = [perm.role_id for perm in permissions]
            if permission_role_ids:
                role = db.query(Role).filter(Role.id == permission_role_ids[0]).first()
                knowledge_name = f"knowledge_{role.role_code}" if role and role.role_code else "knowledge_public"
            else:
                knowledge_name = "knowledge_public"
        
        # 调用文档解析API
        document_upload_url = "http://10.168.27.191:8888/document/upload/url"
        if document.file_url:
            logger.info(f"开始解析文档: {document.filename}, URL: {document.file_url}, 知识库: {knowledge_name}")
            try:
                async with httpx.AsyncClient() as client:
                    upload_response = await client.post(
                        document_upload_url,
                        data={
                            "url": document.file_url,
                            "knowledge_name": knowledge_name
                        }
                    )
                    upload_response.raise_for_status()
                    upload_result = upload_response.json()
                    track_id = upload_result.get("track_id")
                    message = upload_result.get("message", "")
                    logger.info(f"文档解析请求成功: {message}, track_id: {track_id}")
                    
                    # 记录trace_id到数据库
                    if track_id:
                        document.trace_id = track_id
                        db.commit()
                        logger.info(f"已记录trace_id到数据库: {track_id} (文档ID: {document_id})")
                    
                    # 如果文档已存在且状态为completed，直接标记为已处理
                    if "already exists" in message.lower() and "completed" in message.lower():
                        document.is_processed = True
                        db.commit()
                        logger.info(f"文档已存在且已完成，直接标记为已处理: {document.filename} (ID: {document_id})")
                    elif track_id:
                        # 添加后台任务轮询解析进度
                        background_tasks.add_task(
                            poll_document_status,
                            track_id=track_id,
                            knowledge_name=knowledge_name,
                            document_id=document_id,
                            document_type=document_type
                        )
                    else:
                        # track_id为None且不是已存在状态，报错
                        error_msg = f"文档解析失败: 未获取到有效的track_id (消息: {message})"
                        logger.error(error_msg)
                        raise HTTPException(
                            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                            detail=error_msg
                        )
            except Exception as e:
                logger.error(f"文档解析失败: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"文档解析失败: {str(e)}"
                )
        else:
            logger.warning(f"文档没有file_url，跳过解析: {document.filename}")
        
        logger.info(f"文档处理已启动: {document.filename} (ID: {document_id})")
        
        return {
            "success": True,
            "message": "文档处理已启动",
            "document_id": document_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"处理文档失败: {str(e)} (用户ID={current_user.id})")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"处理文档失败: {str(e)}"
        )